from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading('[Your Name]', 0)
doc.add_paragraph('Remote - [Email] - [Phone] - [LinkedIn/Portfolio]')
doc.add_paragraph('Location-independent | SCUBA Instructor | Remote Ops & Content Specialist')

# Profile Section
doc.add_heading('Profile', level=1)
doc.add_paragraph(
    "Globetrotting SCUBA instructor with a background in remote content writing, web development, and customer service. "
    "Equally comfortable underwater or online. Experienced in managing teams, solving problems on the fly, and delivering value—"
    "whether it's a dive briefing or a product launch. Built for travel, remote work, and unconventional workflows."
)

# Remote-Friendly Experience
doc.add_heading('Remote-Friendly Experience', level=1)

doc.add_paragraph('Freelance Full Stack Developer — Remote (Apr 2014 – Jul 2018)', style='List Bullet')
doc.add_paragraph('- Built and maintained websites and blogs for small businesses')
doc.add_paragraph('- Managed social media accounts and wrote content')
doc.add_paragraph('- Developed web apps with clean, functional design')

doc.add_paragraph('Prime Focus Technologies — Remote Content Writer (Apr 2018 – Aug 2018)', style='List Bullet')
doc.add_paragraph('- Wrote compelling show synopses for Netflix and Hotstar')
doc.add_paragraph('- Delivered high-volume content on tight deadlines')
doc.add_paragraph('- Collaborated virtually with editors and QA teams')

doc.add_paragraph('Amazon — Customer Service Associate (WFH Model) (Sep 2020 – Jan 2021)', style='List Bullet')
doc.add_paragraph('- Resolved global customer queries under pressure')
doc.add_paragraph('- Used internal tools to troubleshoot and communicate solutions')
doc.add_paragraph('- Worked across time zones and teams')

# Adventure Ops & Dive Work
doc.add_heading('Adventure Ops & Dive Work', level=1)

doc.add_paragraph('Flying Fish SCUBA, Goa — OWSI (Dec 2022 – Nov 2024)', style='List Bullet')
doc.add_paragraph('Bond Safari, Kovalam — OWSI (Dec 2021 – Dec 2022)', style='List Bullet')
doc.add_paragraph('West Coast Adventures, Kaup/Murdeshwar — OWSI (Oct 2018 – Sep 2020, Feb 2021 – Nov 2021)', style='List Bullet')
doc.add_paragraph('- Taught beginner to advanced dive courses')
doc.add_paragraph('- Managed bookings, equipment, and daily ops')
doc.add_paragraph('- Handled shop finances, budgeting, and team logistics')
doc.add_paragraph('- Delivered 5-star guest experiences in remote coastal locations')
doc.add_paragraph('- Balanced adventure with responsibility in high-pressure settings')

# Skills
doc.add_heading('Skills', level=1)
skills = [
    "Remote work & team coordination",
    "SCUBA instruction (PADI OWSI)",
    "Web dev: HTML, CSS, JS",
    "Content creation & storytelling",
    "Customer support systems (Zendesk, CRM tools)",
    "Budgeting, logistics, and operations",
    "Adaptability, travel readiness, self-motivation"
]
for skill in skills:
    doc.add_paragraph(f'- {skill}', style='List Bullet')

# Certifications
doc.add_heading('Certifications', level=1)
certs = [
    "PADI Open Water Scuba Instructor (OWSI)",
    "Emergency First Response (EFR)",
    "Remote First Aid (Wilderness/Marine preferred - list if any)"
]
for cert in certs:
    doc.add_paragraph(f'- {cert}', style='List Bullet')

# Save the document
docx_path = "Digital_Nomad_Resume.docx"
doc.save(docx_path)

docx_path
